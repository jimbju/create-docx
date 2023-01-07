from docx import Document

# https://python-docx.readthedocs.io/en/latest/user/quickstart.html
document = Document()

document.add_heading('Heading 1')
document.add_heading('Heading 2', level=2)

table = document.add_table(rows=2, cols=2)
row = table.rows[1]
row.cells[0].text = 'Foo bar to you.'
row.cells[1].text = 'And a hearty foo bar to you too sir!'

# Iterate each row in a table
for row in table.rows:
    for cell in row.cells:
        print(cell.text)

# Add row to table
row = table.add_row()

# Add image
#document.add_picture('image-filename.png')

document.save('test.docx')